#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Nov  7 15:14:50 2016

@author: ignacio
"""
import json
import requests
from bs4 import BeautifulSoup
import os
import pandas as pd
import names
import romkan

last_number = '587' #change the number

def requests_data(n):
    # requests bansuke json from website
    # n = 1: 幕內,  n = 2: 十兩
    n = str(n)
    domain = 'http://www.sumo.or.jp'
    url = 'http://www.sumo.or.jp/ResultBanzuke/table_ajax/{}/1'.format(n) #/1/1 /2/1
    payload = {'kakuzuke_id': n, #1 or 2
               'basho_id':last_number, 
               'page':'1'}
    headers = {'X-Requested-With': 'XMLHttpRequest'}
    r = requests.post(url, data = payload, headers=headers)
    r = json.loads(r.text)
    return r

def extract_people(r):
    # contest table
    r = r['BanzukeTable']
    contest = []
    for n in range(len(r)):
        if n % 2 == 0:
            people_east = dict()
            shikona = r[n]['shikona'].split()
            if not shikona:
                shikona = ['','']
            people_east['text'] = shikona + [r[n]['pref_name'], '/', r[n]['heya_name']]
            people_east['img'] = r[n]['photo']
            banzuke = r[n]['banzuke_name']
        else:
            people_west = dict()
            shikona = r[n]['shikona'].split()
            if not shikona:
                shikona = ['','']
            people_west['text'] = shikona + [r[n]['pref_name'], '/', r[n]['heya_name']]
            people_west['img'] = r[n]['photo']
            contest.append([people_east, banzuke, people_west])

    return contest


def update_image(contest):
    images = os.listdir("img")

    # Downloads the images
    for row in contest:
        for i in [0,2]: 
            url = "http://sumo.or.jp/img/sumo_data/rikishi/60x60/" + row[i]['img']
            # url = domain + row[i]['img']
            name = os.path.basename(url)
            if name and name not in images: # check if already downloaded
                response = requests.get(url)
                if response.status_code == 200:
                    f = open("img/" + name, 'wb')
                    f.write(response.content)
                    f.close()

def update_names(contest):
    df = pd.read_csv('names.txt')
    file_names = open("names.txt",'a')

    names_dict = {row[0]:row[1] for _,row in df.iterrows()}

    for row in contest:
        for i in [0,2]:
            try:
                name = row[i]['text'][0]+row[i]['text'][1]
                if name and name not in names_dict.keys():
                    hirakana = names.kanji2hirakana(name)
                    file_names.write(name+","+hirakana+"\n")
                    print("update names:{},{}".format(name, hirakana))
                    names_dict[name] = hirakana
            except:
                pass
    file_names.close()

def write_docx(contest, month, banzuke_type):
    # Write in the docx
    from docx import Document
    from docx.shared import Inches,Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    document.add_heading('番付表', 0)
    paragraph = document.add_paragraph(month)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_heading(banzuke_type, level=1)

    table = document.add_table(1, 5)
    table.style = 'TableGrid'
    table.autofit = False
    table.rows[0].cells[0].text = '東'
    table.rows[0].cells[2].text = '番付'
    table.rows[0].cells[3].text = '西'

    df = pd.read_csv('names.txt')
    names_dict = {row[0]:row[1] for _,row in df.iterrows()}

    for row in contest:
        row_cells = table.add_row().cells
        for j in [0,1]:
            img_name = os.path.basename(row[j*2]['img']) # j = 0,2
            text = ''
            if img_name and img_name != "dummy.gif":
                run = row_cells[j*3].paragraphs[0].add_run()
                try:
                    run.add_picture('img/' + img_name, width = 600000, height = 600000)
                except:
                    print(img_name)
                description = list(row[j*2]['text'])# j = 0,2
                name = description[0]+description[1]
                hirakana = names_dict[name]
                romaji = romkan.to_roma(hirakana)
                description.insert(2,hirakana)
                description.insert(3,romaji)
                text = '{0} {1}\n{2}\n{3}\n{4}{5}{6}'.format(*description)

            row_cells[j*3+1].text = text # j = 1,4
        row_cells[2].text = row[1]

    # set the table width
    table.autofit = False
    for i in [0,3]:
        table.columns[i].width = 650000
        for cell in table.column_cells(i):
            cell.width = 650000
    table.columns[2].width = 500000
    for cell in table.column_cells(2):
        cell.width = 500000
    for i in [1,4]:
        table.columns[i].width = 1800000
        for cell in table.column_cells(i):
            cell.width = 1800000

    document.save('docx_test/'+month+'_'+banzuke_type+'.docx')

if __name__ == "__main__":
    r = requests_data(1)
    month = r['year_jp'] + r['basho_name']
    banzuke_type = r['Kakuzuke']
    contest = extract_people(r)
    update_image(contest)
    update_names(contest)
    write_docx(contest, month, banzuke_type)
